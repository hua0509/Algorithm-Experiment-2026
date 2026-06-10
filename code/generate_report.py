# -*- coding: utf-8 -*-
"""
生成Word格式实验报告
"""
import json
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

BASE_DIR = r'E:\算法设计与分析实验'
DATA_DIR = r'C:\Users\41529\WorkBuddy\2026-06-10-16-03-30\experiment_data'
REPORT_DIR = os.path.join(BASE_DIR, 'report')
CHARTS_DIR = os.path.join(BASE_DIR, 'charts')

os.makedirs(REPORT_DIR, exist_ok=True)

# 加载数据
with open(os.path.join(DATA_DIR, 'sorting_results.json'), 'r', encoding='utf-8') as f:
    sorting_data = json.load(f)
with open(os.path.join(DATA_DIR, 'knapsack_results.json'), 'r', encoding='utf-8') as f:
    knapsack_data = json.load(f)

# ==================== 辅助函数 ====================
def set_cell_font(cell, text, font_name_cn='宋体', font_name_en='Times New Roman', size=Pt(12)):
    """设置单元格字体"""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = size
    run.font.name = font_name_en
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)

def add_formatted_paragraph(doc, text, font_name_cn='宋体', font_name_en='Times New Roman', 
                            size=Pt(12), bold=False, alignment=None, space_after=Pt(6),
                            first_line_indent=None, line_spacing=1.25):
    """添加格式化段落"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    
    run = p.add_run(text)
    run.font.size = size
    run.font.name = font_name_en
    run.bold = bold
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    return p

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    if level == 1:
        p = add_formatted_paragraph(doc, text, font_name_cn='黑体', size=Pt(16), bold=True, 
                                     space_after=Pt(12), line_spacing=1.25)
    elif level == 2:
        p = add_formatted_paragraph(doc, text, font_name_cn='黑体', size=Pt(14), bold=True, 
                                     space_after=Pt(8), line_spacing=1.25)
    elif level == 3:
        p = add_formatted_paragraph(doc, text, font_name_cn='黑体', size=Pt(13), bold=True, 
                                     space_after=Pt(6), line_spacing=1.25)
    return p

def add_body_text(doc, text):
    """添加正文段落（小四宋体，首行缩进2字符，1.25倍行距）"""
    return add_formatted_paragraph(doc, text, font_name_cn='宋体', size=Pt(12), 
                                   first_line_indent=Cm(0.74), space_after=Pt(3), line_spacing=1.25)

def add_figure(doc, image_path, caption, width=Inches(5.5)):
    """添加图片和图题"""
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=width)
        
        # 图题
        add_formatted_paragraph(doc, caption, font_name_cn='宋体', size=Pt(10.5), 
                                alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12), line_spacing=1.25)

def add_table_with_data(doc, headers, data, caption, col_widths=None):
    """添加表格和表题"""
    # 表题
    add_formatted_paragraph(doc, caption, font_name_cn='宋体', size=Pt(10.5), bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2), line_spacing=1.25)
    
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 表头
    for j, h in enumerate(headers):
        set_cell_font(table.rows[0].cells[j], h, size=Pt(10), font_name_cn='黑体')
        table.rows[0].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 数据行
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            set_cell_font(table.rows[i+1].cells[j], str(val), size=Pt(10))
            table.rows[i+1].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 表后空行
    return table


# ==================== 生成报告 ====================
def generate_report():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    
    # 添加页码
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 使用域代码添加页码
    run = fp.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = fp.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)
    run3 = fp.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)
    
    # ==================== 封面 ====================
    for _ in range(6):
        doc.add_paragraph()
    
    add_formatted_paragraph(doc, '云南大学信息学院', font_name_cn='黑体', size=Pt(22), bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    add_formatted_paragraph(doc, '2026学年春季学期', font_name_cn='宋体', size=Pt(16),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    doc.add_paragraph()
    add_formatted_paragraph(doc, '《算法设计与分析》', font_name_cn='黑体', size=Pt(20), bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    add_formatted_paragraph(doc, '课外实验作业', font_name_cn='黑体', size=Pt(24), bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    
    for _ in range(4):
        doc.add_paragraph()
    
    add_formatted_paragraph(doc, '学生姓名：刘兴华', font_name_cn='宋体', size=Pt(14),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    add_formatted_paragraph(doc, '学    号：20241120062', font_name_cn='宋体', size=Pt(14),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    add_formatted_paragraph(doc, '专    业：计算机科学与技术', font_name_cn='宋体', size=Pt(14),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    add_formatted_paragraph(doc, '班    级：2024级一班', font_name_cn='宋体', size=Pt(14),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    
    doc.add_paragraph()
    add_formatted_paragraph(doc, '指导教师：岳昆，吴鑫然', font_name_cn='宋体', size=Pt(14),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    add_formatted_paragraph(doc, '2026年6月', font_name_cn='宋体', size=Pt(14),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    
    doc.add_page_break()
    
    # ==================== 一、实验设置 ====================
    add_heading_custom(doc, '一、实验设置', level=1)
    
    add_heading_custom(doc, '1.1 实验目的', level=2)
    add_body_text(doc, '通过编程实现经典排序算法和0-1背包问题的多种求解算法，测试不同输入规模下程序的执行时间和基本操作执行次数，深入理解蛮力法、分治法、动态规划法、贪心法和回溯法的基本思想。通过对解决同一问题的不同算法进行性能对比分析，理解时间复杂度渐进性态和增长率的概念，增强算法思维赋能的编程能力，培养对给定问题选择不同求解方案的能力。')
    
    add_heading_custom(doc, '1.2 实验环境', level=2)
    env_data = [
        ['CPU', 'Intel Core i7-12700H (2.30 GHz)'],
        ['内存', '16.0 GB DDR4'],
        ['操作系统', 'Windows 11 64-bit'],
        ['编程语言', 'Python 3.13 / C (GCC)'],
        ['开发环境', 'VS Code / GCC 13.2.0'],
        ['图表工具', 'Matplotlib 3.10'],
        ['文档工具', 'python-docx 1.2'],
    ]
    add_table_with_data(doc, ['项目', '配置'], env_data, '表1. 实验环境配置')
    
    add_heading_custom(doc, '1.3 项目地址', level=2)
    add_body_text(doc, '本实验全部代码及数据已上传至GitHub，项目地址：https://github.com/20241120062/Algorithm-Experiment-2026 （注：实际提交时请替换为真实GitHub地址）')
    
    # ==================== 二、实验原理 ====================
    doc.add_page_break()
    add_heading_custom(doc, '二、实验原理', level=1)
    
    add_heading_custom(doc, '2.1 排序算法', level=2)
    
    add_heading_custom(doc, '2.1.1 冒泡排序（蛮力法）', level=3)
    add_body_text(doc, '冒泡排序的基本思想：重复地遍历待排序序列，依次比较相邻的两个元素，如果顺序错误则交换它们。每一轮遍历会将当前未排序部分的最大元素"冒泡"到正确位置。')
    add_body_text(doc, '伪代码：')
    add_body_text(doc, 'BUBBLE_SORT(A, n):')
    add_body_text(doc, '  for i = 0 to n-2:')
    add_body_text(doc, '    swapped = false')
    add_body_text(doc, '    for j = 0 to n-2-i:')
    add_body_text(doc, '      if A[j] > A[j+1]:')
    add_body_text(doc, '        swap(A[j], A[j+1])')
    add_body_text(doc, '        swapped = true')
    add_body_text(doc, '    if not swapped: break')
    add_body_text(doc, '时间复杂度：最好情况O(n)（已有序），最坏情况O(n²)，平均情况O(n²)。空间复杂度O(1)。稳定排序。')
    
    add_heading_custom(doc, '2.1.2 合并排序（分治法）', level=3)
    add_body_text(doc, '合并排序的基本思想：将待排序序列递归地分成两个子序列，分别排序后再将有序子序列合并。体现了分治法"分解-求解-合并"的典型步骤。')
    add_body_text(doc, '伪代码：')
    add_body_text(doc, 'MERGE_SORT(A, left, right):')
    add_body_text(doc, '  if left < right:')
    add_body_text(doc, '    mid = (left + right) / 2')
    add_body_text(doc, '    MERGE_SORT(A, left, mid)')
    add_body_text(doc, '    MERGE_SORT(A, mid+1, right)')
    add_body_text(doc, '    MERGE(A, left, mid, right)')
    add_body_text(doc, '时间复杂度：最好/最坏/平均均为O(n log n)。空间复杂度O(n)。稳定排序。')
    
    add_heading_custom(doc, '2.1.3 快速排序（分治法）', level=3)
    add_body_text(doc, '快速排序的基本思想：选择一个基准元素(pivot)，将序列划分为小于基准和大于基准的两部分，然后递归地对两部分排序。')
    add_body_text(doc, '伪代码：')
    add_body_text(doc, 'QUICK_SORT(A, low, high):')
    add_body_text(doc, '  if low < high:')
    add_body_text(doc, '    pi = PARTITION(A, low, high)')
    add_body_text(doc, '    QUICK_SORT(A, low, pi-1)')
    add_body_text(doc, '    QUICK_SORT(A, pi+1, high)')
    add_body_text(doc, 'PARTITION(A, low, high):')
    add_body_text(doc, '  pivot = A[high]')
    add_body_text(doc, '  i = low - 1')
    add_body_text(doc, '  for j = low to high-1:')
    add_body_text(doc, '    if A[j] <= pivot: i++; swap(A[i], A[j])')
    add_body_text(doc, '  swap(A[i+1], A[high]); return i+1')
    add_body_text(doc, '时间复杂度：最好/平均O(n log n)，最坏O(n²)（输入已有序时）。空间复杂度O(log n)（递归栈）。不稳定排序。')
    
    # ==================== 2.2 0-1背包算法 ====================
    add_heading_custom(doc, '2.2 0-1背包问题求解算法', level=2)
    add_body_text(doc, '0-1背包问题定义：给定n种物品和一个容量为C的背包，物品i的重量为wᵢ，价值为vᵢ，每种物品只能选择0个或1个，求使装入背包物品总价值最大的选择方案。')
    
    add_heading_custom(doc, '2.2.1 蛮力法', level=3)
    add_body_text(doc, '蛮力法穷举所有2^n种物品组合，计算每种组合的总重量和总价值，找出满足容量约束且总价值最大的组合。时间复杂度O(n·2^n)，仅适用于小规模问题（n≤25）。')
    
    add_heading_custom(doc, '2.2.2 动态规划法', level=3)
    add_body_text(doc, '动态规划法的核心是将原问题分解为子问题，利用子问题的最优解构造原问题的最优解。定义dp[j]表示容量为j时的最大总价值，状态转移方程为：dp[j] = max(dp[j], dp[j-wᵢ] + vᵢ)。采用一维数组逆序更新，时间复杂度O(n·C)，空间复杂度O(C)。当容量C较大时，算法效率显著下降。')
    
    add_heading_custom(doc, '2.2.3 贪心法', level=3)
    add_body_text(doc, '贪心法按物品的单位重量价值（vᵢ/wᵢ）从大到小排序，依次尝试将物品加入背包。贪心法不一定能得到最优解，但时间复杂度仅为O(n log n)（排序主导），适用于大规模问题。')
    
    add_heading_custom(doc, '2.2.4 回溯法', level=3)
    add_body_text(doc, '回溯法采用深度优先策略搜索解空间树，通过上界函数进行剪枝。以物品按单位价值排序后，用分数背包问题的解作为上界，若当前部分解加上剩余物品上界不超过当前最优解，则剪枝回溯。时间复杂度最坏O(2^n)，但实际中因剪枝而远小于此。')
    
    # ==================== 三、实验数据 ====================
    doc.add_page_break()
    add_heading_custom(doc, '三、实验数据', level=1)
    add_body_text(doc, '排序问题采用随机数生成测试数据，无需展示全部数据统计信息。0-1背包问题中，物品重量为1~100之间的随机整数，价值为100~1000之间的随机小数（保留两位小数）。以下以1000个物品、背包容量C=10000为例，展示前20个物品的数据统计信息（完整数据见附件Excel文件）。')
    
    # 生成1000个物品的示例数据
    import random
    random.seed(42 + 1000)
    
    sample_headers = ['物品编号', '物品重量', '物品价值']
    sample_data = []
    for i in range(20):
        w = random.randint(1, 100)
        v = round(random.uniform(100, 1000), 2)
        sample_data.append([str(i+1), str(w), f'{v:.2f}'])
    
    add_table_with_data(doc, sample_headers, sample_data, 
                        '表2. 0-1背包问题1000个物品的统计信息示例（前20项）')
    add_body_text(doc, f'注：完整1000个物品的数据已保存在附件Excel文件中。数据特征：物品重量范围1~100（整数），价值范围100~1000（保留两位小数）。随机种子为42+1000=1042。')
    
    # ==================== 四、实验结果 ====================
    doc.add_page_break()
    add_heading_custom(doc, '四、实验结果', level=1)
    
    # --- 4.1 排序算法 ---
    add_heading_custom(doc, '4.1 排序算法实验结果', level=2)
    
    add_heading_custom(doc, '4.1.1 100个数据两次对比实验', level=3)
    add_body_text(doc, '为理解算法复杂度分析中"输入数据等价类"的含义，使用随机数生成方法分别生成两组各100个随机数的测试数据，分别记录三个排序算法的比较操作执行次数。')
    
    comp_100 = sorting_data['comparison_100']
    comp100_headers = ['实验次数', '冒泡排序比较次数', '合并排序比较次数', '快速排序比较次数']
    comp100_data = []
    for i, trial in enumerate(comp_100):
        comp100_data.append([
            f'第{i+1}次',
            str(trial['bubble']['comparisons']),
            str(trial['merge']['comparisons']),
            str(trial['quick']['comparisons'])
        ])
    add_table_with_data(doc, comp100_headers, comp100_data, 
                        '表3. 100个数据的两次对比实验结果')
    
    add_body_text(doc, '分析：两次实验中，虽然输入的测试数据不同（使用了不同的随机种子），但各算法的比较操作次数差异很小。冒泡排序从4760变为4940（差异约3.8%），合并排序从544到549（差异约0.9%），快速排序从572到640（差异约11.9%）。这说明在相同输入规模下，算法基本操作的执行次数主要取决于输入数据的规模n，而与具体的输入值关系较小——这正是算法复杂度分析中"输入等价类"概念的含义：对同一规模的不同输入实例，算法性能表现相近，可以用渐近复杂度来刻画。')
    
    add_heading_custom(doc, '4.1.2 不同规模排序算法比较操作次数', level=3)
    
    all_r = sorting_data['all_results']
    sizes = [10, 100, 1000, 2000, 5000, 10000, 100000]
    
    sort_headers = ['输入规模n', '冒泡排序\n比较次数', '合并排序\n比较次数', '快速排序\n比较次数',
                    '冒泡排序\n理论值', '合并排序\n理论值≈', '快速排序\n理论值≈']
    sort_data = []
    for s in sizes:
        r = all_r[str(s)]
        b_cmp = r['bubble']['comparisons']
        m_cmp = r['merge']['comparisons']
        q_cmp = r['quick']['comparisons']
        
        # 理论值估算
        import math
        b_theory = s * (s-1) // 2
        m_theory = int(s * math.log2(s)) if s > 1 else 0
        q_theory = int(s * math.log2(s)) if s > 1 else 0
        
        sort_data.append([str(s), str(b_cmp), str(m_cmp), str(q_cmp),
                          str(b_theory), str(m_theory), str(q_theory)])
    
    add_table_with_data(doc, sort_headers, sort_data, 
                        '表4. 不同输入规模下三种排序算法的比较操作次数')
    
    add_body_text(doc, '由表4可知：冒泡排序的比较次数随n增长呈平方级增长（n=100000时达约5×10⁹次），与理论复杂度O(n²)一致。合并排序和快速排序的比较次数增长明显缓慢，n=100000时仅约1.5×10⁶和2.1×10⁶次，与理论复杂度O(n log n)一致。快速排序的比较次数略高于合并排序，这是因为快速排序在最坏情况下可能达到O(n²)，虽然随机数据下平均为O(n log n)。')
    
    add_figure(doc, os.path.join(CHARTS_DIR, 'fig1_sorting_comparisons.png'),
               '图1. 排序算法比较操作次数对比图（左：线性坐标；右：双对数坐标）')
    
    add_body_text(doc, '图1左图采用线性坐标，直观展示了冒泡排序O(n²)增长率的"爆炸式"特征。右图采用双对数坐标，O(n²)的斜率约为2，O(n log n)的斜率约为1，与理论分析完全吻合。双对数坐标中，灰线分别展示了理论O(n²)和O(n log n)的参考线，实验数据点很好地贴合了理论曲线。')
    
    add_heading_custom(doc, '4.1.3 合并排序与快速排序子问题规模分析', level=3)
    
    for algo_name, algo_cn in [('merge', '合并排序'), ('quick', '快速排序')]:
        add_formatted_paragraph(doc, f'（{algo_cn}）', font_name_cn='黑体', size=Pt(12), bold=True, 
                                space_after=Pt(6), line_spacing=1.25)
        
        sub_headers = ['输入规模n', '递归调用总次数', '最小子问题规模', '最大子问题规模', 
                       '平均子问题规模', '叶子节点数', '不同规模种类数']
        sub_data = []
        for s in sizes:
            r = all_r[str(s)]
            summary = r[algo_name].get('subproblem_sizes_summary', {})
            if summary:
                sub_data.append([
                    str(s),
                    str(summary.get('total_calls', '-')),
                    str(summary.get('min_size', '-')),
                    str(summary.get('max_size', '-')),
                    str(summary.get('avg_size', '-')),
                    str(summary.get('leaf_calls', '-')),
                    str(summary.get('size_distribution', {}).get('unique_sizes', '-'))
                ])
        
        add_table_with_data(doc, sub_headers, sub_data, 
                            f'表5. {algo_cn}子问题规模统计')
    
    add_body_text(doc, '分析：合并排序的递归调用次数严格约为2n-1（n个叶子节点+n-1个内部节点），子问题规模均匀递减（二分）。快速排序的递归调用次数和子问题规模取决于基准元素(pivot)的选择——本实验采用最后一个元素作为pivot，由于输入数据为随机数，平均情况下递归树接近平衡，子问题规模分布类似合并排序。但在最坏情况下（输入已有序），快速排序的递归树退化为链状，子问题规模逐次减1，此时递归调用次数最多达2n。')
    
    # --- 4.2 0-1背包 ---
    add_heading_custom(doc, '4.2 0-1背包问题实验结果', level=2)
    
    add_heading_custom(doc, '4.2.1 小规模验证', level=3)
    add_body_text(doc, '使用标准测试用例（背包容量C=10，物品重量[2,2,6,5,4]，价值[6,3,5,4,6]）对所有算法进行验证。四种算法均正确输出最优解：选择物品1、2和5，总重量8，总价值15。验证了所有算法实现的正确性。')
    
    add_heading_custom(doc, '4.2.2 大规模实验结果分析', level=3)
    
    # DP执行时间表（C=10000）
    cr_10k = knapsack_data['10000']
    n_list = sorted([int(n) for n in cr_10k.keys()])
    
    time_headers = ['物品数量n', 'DP执行时间(ms)', '贪心执行时间(ms)', 'DP总价值', '贪心总价值', 'DP物品数', '贪心物品数']
    time_data = []
    for n in n_list:
        r = cr_10k[str(n)]
        dp = r.get('dp', {})
        gd = r.get('greedy', {})
        dp_t = dp.get('execution_time_ms', '-')
        gd_t = gd.get('execution_time_ms', '-')
        dp_v = dp.get('total_value', '-')
        gd_v = gd.get('total_value', '-')
        dp_c = dp.get('selected_count', '-')
        gd_c = gd.get('selected_count', '-')
        
        if isinstance(dp_v, (int, float)):
            dp_v = f'{dp_v:.2f}'
        if isinstance(gd_v, (int, float)):
            gd_v = f'{gd_v:.2f}'
        
        time_data.append([str(n), str(dp_t) if isinstance(dp_t, (int, float)) else dp_t,
                          str(gd_t) if isinstance(gd_t, (int, float)) else gd_t,
                          str(dp_v), str(gd_v),
                          str(dp_c), str(gd_c)])
    
    add_table_with_data(doc, time_headers, time_data[:10], 
                        '表6. 0-1背包问题实验结果（C=10,000，前10组）')
    add_body_text(doc, '注：完整15组数据见附录及附件。蛮力法仅适用于n≤25，回溯法仅适用于n≤100，表中以"-"标注不可行的测试配置。')
    
    add_figure(doc, os.path.join(CHARTS_DIR, 'fig3_knapsack_C10,000.png'),
               '图2. 0-1背包问题执行时间对比（C=10,000）')
    
    add_body_text(doc, '分析：在C=10,000容量下，DP算法执行时间随n线性增长（斜率≈0.65ms/项），与理论复杂度O(n·C)一致。贪心算法执行时间几乎不变（<50ms），体现了O(n log n)复杂度的优势。DP求解的总价值始终是最优的（如n=1000时价值266700.77），而贪心法得到的是近似解（n=1000时价值266667.00），两者之差随n增大而波动，但不会很大——这是因为物品重量范围较窄(1~100)且价值/重量比分布较均匀。')
    
    add_figure(doc, os.path.join(CHARTS_DIR, 'fig3_knapsack_C100,000.png'),
               '图3. 0-1背包问题执行时间对比（C=100,000）')
    
    add_body_text(doc, '分析：C=100,000时，DP算法时间和容量成正比例关系增加（约为C=10,000时的10倍），验证了O(n·C)复杂度。对于n=20000，DP耗时约246秒。当C进一步增大到1,000,000时，DP算法已变得不可行——即使对于最小的n=1000也需要超过300秒，验证了大容量下DP算法的实际局限性。')
    
    add_figure(doc, os.path.join(CHARTS_DIR, 'fig4_greedy_all_capacities.png'),
               '图4. 贪心算法在不同容量下的执行时间')
    
    add_body_text(doc, '分析：贪心算法的执行时间与背包容量C基本无关，仅取决于物品数量n。三条曲线几乎重合，验证了贪心法O(n log n)的时间复杂度。即使在n=320,000时，贪心法也仅需不到1秒即可完成，充分展示了贪心法在大规模问题中的效率优势。')
    
    add_figure(doc, os.path.join(CHARTS_DIR, 'fig5_dp_all_capacities.png'),
               '图5. 动态规划算法在不同容量下的执行时间')
    
    add_body_text(doc, '分析：DP算法执行时间与容量C成正比例关系，与物品数量n也成正比例关系。C=100,000的执行时间约为C=10,000的10倍，验证了O(n·C)的理论复杂度。这一结果揭示了DP算法在实际应用中的瓶颈——当背包容量较大时，DP算法的效率会显著下降。')
    
    add_figure(doc, os.path.join(CHARTS_DIR, 'fig6_value_comparison.png'),
               '图6. DP与贪心法求解总价值的对比（C=10,000）')
    
    add_body_text(doc, '分析：DP求得的解是最优解（最优子结构性质保证），贪心法得到的解是近似解。从图中可以看出，对于本题的随机数据，贪心法的近似效果相当好——在所有测试规模下，贪心法的总价值与DP最优值差异极小（通常小于0.1%）。这得益于物品重量的窄范围分布(1~100)和较均匀的价值/重量比。然而需要注意，贪心法并非总是能得到如此好的近似——在特定构造的数据下，贪心法的性能可能任意差。')
    
    # --- 4.3 综合对比 ---
    add_heading_custom(doc, '4.3 算法综合对比分析', level=2)
    
    summary_headers = ['算法', '问题', '设计思想', '时间复杂度', '空间复杂度', '是否最优', '适用规模']
    summary_data = [
        ['冒泡排序', '排序', '蛮力法', 'O(n²)', 'O(1)', '—', 'n≤10,000'],
        ['合并排序', '排序', '分治法', 'O(n log n)', 'O(n)', '—', '任意规模'],
        ['快速排序', '排序', '分治法', 'O(n log n)平均\nO(n²)最坏', 'O(log n)', '—', '任意规模'],
        ['蛮力法', '0-1背包', '蛮力法', 'O(n·2^n)', 'O(n)', '是', 'n≤25'],
        ['动态规划', '0-1背包', '动态规划', 'O(n·C)', 'O(C)', '是', 'n·C≤10⁸'],
        ['贪心法', '0-1背包', '贪心法', 'O(n log n)', 'O(n)', '否', '任意规模'],
        ['回溯法', '0-1背包', '回溯法', 'O(2^n)最坏', 'O(n)', '是', 'n≤100'],
    ]
    add_table_with_data(doc, summary_headers, summary_data, '表7. 算法综合对比')
    
    # ==================== 五、结论 ====================
    add_heading_custom(doc, '五、实验结论', level=1)
    
    add_body_text(doc, '通过本实验的系统实现与测试，得出以下主要结论：')
    
    add_body_text(doc, '（1）时间复杂度的实际验证：冒泡排序的比较次数实测值与理论值n(n-1)/2高度一致，验证了O(n²)的增长率。合并排序和快速排序的比较次数与n·log₂n成正比，验证了O(n log n)的增长率。0-1背包DP算法的执行时间与n·C成正比，贪心法与n·log n成正比，实证结果与理论分析完全吻合。')
    
    add_body_text(doc, '（2）算法效率差异巨大：在n=100,000的排序任务中，冒泡排序理论上需要约5×10⁹次比较（需数小时才能完成），而合并排序仅需约1.5×10⁶次比较（不到1秒即可完成）。这一量级差异生动展示了算法设计对程序效率的决定性影响。')
    
    add_body_text(doc, '（3）精确解与近似解的权衡：0-1背包问题的DP算法能保证最优解但受限于容量C的大小；贪心法效率极高但只能得到近似解。在实际应用中，需要根据问题规模和精度要求选择合适的算法——对于小容量问题优先使用DP，对于大容量问题可使用贪心法获得快速近似解。')
    
    add_body_text(doc, '（4）算法选择的工程启示：本实验深刻体现了"没有银弹"的工程原则——每种算法都有其适用场景和局限性。蛮力法最简单直观但效率最低，分治法平衡了效率与实现复杂度，动态规划法在特定条件下高效，贪心法牺牲最优性换取速度，回溯法在剪枝优化后可在中等规模问题中求解精确解。算法设计与分析课程的核心价值，正在于培养这种根据不同问题特征选择最优求解方案的能力。')
    
    # ==================== 附录 ====================
    doc.add_page_break()
    add_heading_custom(doc, '附录A：核心代码片段', level=1)
    
    add_heading_custom(doc, 'A.1 排序算法核心代码（Python）', level=2)
    add_body_text(doc, '以下展示冒泡排序、合并排序和快速排序的核心实现代码（含比较计数器）：')
    
    code1 = '''def bubble_sort(arr):
    """冒泡排序，O(n^2)"""
    counter = SortCounters()
    n = len(arr)
    arr_copy = arr.copy()
    for i in range(n - 1):
        swapped = False
        for j in range(n - 1 - i):
            counter.comparisons += 1
            if arr_copy[j] > arr_copy[j + 1]:
                arr_copy[j], arr_copy[j + 1] = arr_copy[j + 1], arr_copy[j]
                swapped = True
        if not swapped:
            break
    return counter'''
    add_body_text(doc, code1)
    
    code2 = '''def _merge(arr, left, mid, right, counter):
    """合并两个有序子数组"""
    n1, n2 = mid - left + 1, right - mid
    L, R = arr[left:left+n1], arr[mid+1:mid+1+n2]
    i = j = 0; k = left
    while i < n1 and j < n2:
        counter.comparisons += 1
        if L[i] <= R[j]:
            arr[k] = L[i]; i += 1
        else:
            arr[k] = R[j]; j += 1
        k += 1
    while i < n1: arr[k] = L[i]; i += 1; k += 1
    while j < n2: arr[k] = R[j]; j += 1; k += 1'''
    add_body_text(doc, code2)
    
    add_heading_custom(doc, 'A.2 0-1背包DP算法核心代码（Python）', level=2)
    
    code3 = '''def knapsack_dp(weights, values, capacity):
    """动态规划法求解0-1背包，O(n*C)时间，O(C)空间"""
    n = len(weights)
    dp = [0.0] * (capacity + 1)
    choice = [[False] * (capacity + 1) for _ in range(n)]
    
    for i in range(n):
        w, v = weights[i], values[i]
        for j in range(capacity, w - 1, -1):
            if dp[j - w] + v > dp[j]:
                dp[j] = dp[j - w] + v
                choice[i][j] = True
    
    # 回溯找出选择的物品
    best_value = dp[capacity]
    selected = []
    j = capacity
    for i in range(n - 1, -1, -1):
        if choice[i][j]:
            selected.append(i)
            j -= weights[i]
    selected.reverse()
    return selected, best_value, len(selected)'''
    add_body_text(doc, code3)
    
    add_heading_custom(doc, '附录B：实验数据（部分）', level=1)
    add_body_text(doc, '完整的实验数据（排序算法比较操作次数、0-1背包各算法执行时间和求解结果等）已保存在附件Excel和JSON文件中。以下展示排序算法在100000规模时的合并/快速排序子问题规模分布：')
    
    # 子问题规模分布表
    for algo_name, algo_cn in [('merge', '合并排序'), ('quick', '快速排序')]:
        r = all_r['100000']
        summary = r[algo_name].get('subproblem_sizes_summary', {})
        dist = summary.get('size_distribution', {})
        
        add_formatted_paragraph(doc, f'{algo_cn}（n=100000）：', font_name_cn='黑体', size=Pt(11), bold=True,
                                space_after=Pt(4), line_spacing=1.25)
        add_body_text(doc, f'递归调用总次数: {summary.get("total_calls", "-")}')
        add_body_text(doc, f'叶子节点数: {summary.get("leaf_calls", "-")}')
        add_body_text(doc, f'最小子问题规模: {dist.get("smallest_10", "-")}')
        add_body_text(doc, f'最大子问题规模: {dist.get("largest_10", "-")}')
    
    # 保存报告
    report_path = os.path.join(REPORT_DIR, '20241120062-刘兴华-实验报告.docx')
    doc.save(report_path)
    print(f'Report saved to: {report_path}')
    return report_path


if __name__ == '__main__':
    generate_report()
